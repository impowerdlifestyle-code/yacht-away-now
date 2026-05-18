from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

INPUT = "/Users/ciaranwentz/yacht-away-now/partnership-proposal.md"
OUTPUT = "/Users/ciaranwentz/yacht-away-now/Yacht-Away-Now-Partnership-Proposal.docx"

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    hs.font.name = 'Calibri'

with open(INPUT, 'r') as f:
    lines = f.readlines()

def add_bold_runs(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def parse_table_lines(table_lines):
    rows = []
    for line in table_lines:
        line = line.strip().strip('|')
        cells = [c.strip() for c in line.split('|')]
        rows.append(cells)
    if len(rows) >= 2 and all(re.match(r'^[-:]+$', c) for c in rows[1]):
        rows.pop(1)
    return rows

def add_table(rows):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append('')

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            run = p.add_run(clean)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if i == 0:
                run.bold = True
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '1a365d')
                shading.set(qn('w:val'), 'clear')
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                cell._tc.get_or_add_tcPr().append(shading)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')
    stripped = line.strip()

    if not stripped:
        i += 1
        continue

    if stripped == '---':
        doc.add_paragraph('─' * 60)
        i += 1
        continue

    if stripped.startswith('# ') and not stripped.startswith('## '):
        p = doc.add_heading(stripped[2:], level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    if stripped.startswith('#### '):
        doc.add_heading(stripped[5:], level=3)
        i += 1
        continue

    if stripped.startswith('### '):
        doc.add_heading(stripped[4:], level=2)
        i += 1
        continue

    if stripped.startswith('## '):
        doc.add_heading(stripped[3:], level=1)
        i += 1
        continue

    if stripped.startswith('|') and '|' in stripped[1:]:
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        rows = parse_table_lines(table_lines)
        add_table(rows)
        continue

    if stripped.startswith('- [ ] '):
        p = doc.add_paragraph(style='List Bullet')
        add_bold_runs(p, '☐ ' + stripped[6:])
        i += 1
        continue

    if stripped.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        add_bold_runs(p, stripped[2:])
        i += 1
        continue

    if re.match(r'^\d+\.', stripped):
        text = re.sub(r'^\d+\.\s*', '', stripped)
        p = doc.add_paragraph(style='List Number')
        add_bold_runs(p, text)
        i += 1
        continue

    p = doc.add_paragraph()
    add_bold_runs(p, stripped)
    i += 1

doc.save(OUTPUT)
print(f"DOCX generated: {OUTPUT}")
